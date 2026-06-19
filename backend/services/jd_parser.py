from docx import Document
import io
from typing import Optional

def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        # Also pull tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)
        return "\n".join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")